"""
Script untuk generate file template soal .docx
Jalankan: python3 generate_soal.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── Helper: tambah paragraph dengan numbering (list) ───────────────────────

def add_heading_paragraph(doc, text):
    """Bold paragraph sebagai judul seksi"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return p


def add_question(doc, number, text):
    """Paragraf soal dengan penomoran manual: 1. teks"""
    p = doc.add_paragraph(f"{number}. {text}")
    p.runs[0].font.size = Pt(11)
    return p


def add_option(doc, letter, text):
    """Paragraf opsi dengan huruf manual: A. teks"""
    p = doc.add_paragraph(f"{letter}. {text}")
    p.paragraph_format.left_indent = Inches(0.3)
    p.runs[0].font.size = Pt(11)
    return p


def add_meta(doc, kunci, tipe):
    """Paragraf metadata: Kunci: X Tipe: radio"""
    p = doc.add_paragraph(f"Kunci: {kunci} Tipe: {tipe}")
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.size = Pt(10)
    return p


def add_code_block(doc, code_text):
    """Paragraf kode: tiap baris jadi paragraf tersendiri dengan font Courier New"""
    lines = code_text.split('\n')
    paragraphs = []
    for line in lines:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.3)
        paragraphs.append(p)
    return paragraphs


def add_math_block(doc, math_text):
    """Paragraf rumus: tiap baris dengan font Cambria Math (penanda rumus matematika)"""
    lines = math_text.split('\n')
    paragraphs = []
    for line in lines:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = "Cambria Math"
            run.font.size = Pt(12)
        p.paragraph_format.left_indent = Inches(0.3)
        paragraphs.append(p)
    return paragraphs


def add_blank(doc):
    doc.add_paragraph("")

# ─── SOAL 1: Sudah ada – akar persamaan ─────────────────────────────────────
add_question(doc, 1, "Berapakah nilai dari akar persamaan √x + 16 jika x = 9?")
add_option(doc, "A", "17")
add_option(doc, "B", "19")
add_option(doc, "C", "21")
add_option(doc, "D", "25")
add_meta(doc, "B", "radio")
add_blank(doc)

# ─── SOAL 2: Sudah ada – pecahan aljabar ────────────────────────────────────
add_question(doc, 2, "Bentuk sederhana dari pecahan matematika (a² - b²) / (a - b) adalah...")
add_option(doc, "A", "a - b")
add_option(doc, "B", "a + b")
add_option(doc, "C", "a × b")
add_option(doc, "D", "a / b")
add_meta(doc, "B", "radio")
add_blank(doc)

# ─── SOAL 3: Sudah ada – persamaan kuadrat (checkbox) ───────────────────────
add_question(doc, 3, "Manakah dari persamaan berikut yang termasuk persamaan kuadrat? (Pilih semua yang benar)")
add_option(doc, "A", "x² + 2x + 1 = 0")
add_option(doc, "B", "3x + 5 = 11")
add_option(doc, "C", "2x² - 8 = 0")
add_option(doc, "D", "y = 4x³ - 2")
add_meta(doc, "A, C", "checkbox")
add_blank(doc)

# ─── SOAL 4: Sudah ada – luas lingkaran (text) ──────────────────────────────
add_question(doc, 4, "Tuliskan rumus luas lingkaran beserta keterangannya!")
p = doc.add_paragraph("Tipe: text")
p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
p.runs[0].font.size = Pt(10)
add_blank(doc)

# ─── SOAL 5: Sudah ada – gambar Tailwind ────────────────────────────────────
add_question(doc, 5, "Perhatikan Gambar Berikut")
p = doc.add_paragraph("Mana yang Tailwindcss")
p.runs[0].font.size = Pt(11)
add_option(doc, "A", "")
add_option(doc, "B", "Dakjadjk")
add_meta(doc, "A", "radio")
add_blank(doc)

# ════════════════════════════════════════════════════════════════════════════
# ─── SOAL BARU 6: Rumus ABC / Kuadrat ───────────────────────────────────────
# ════════════════════════════════════════════════════════════════════════════
add_question(doc, 6, "Diketahui persamaan kuadrat berikut:")
add_math_block(doc, "2x² - 4x - 6 = 0")
p = doc.add_paragraph("Menggunakan rumus ABC:")
p.runs[0].font.size = Pt(11)
add_math_block(doc, "x = (-b ± √(b² - 4ac)) / 2a")
p = doc.add_paragraph("dengan a = 2, b = -4, c = -6. Berapakah nilai x yang memenuhi?")
p.runs[0].font.size = Pt(11)
add_option(doc, "A", "x = 3 atau x = -1")
add_option(doc, "B", "x = 2 atau x = -1")
add_option(doc, "C", "x = -3 atau x = 1")
add_option(doc, "D", "x = 6 atau x = -2")
add_meta(doc, "A", "radio")
add_blank(doc)

# ─── SOAL 7: Diskriminan ────────────────────────────────────────────────────
add_question(doc, 7, "Pada persamaan kuadrat ax² + bx + c = 0, nilai diskriminan:")
add_math_block(doc, "D = b² - 4ac")
p = doc.add_paragraph("Jika D > 0, maka persamaan memiliki...")
p.runs[0].font.size = Pt(11)
add_option(doc, "A", "Dua akar real berbeda")
add_option(doc, "B", "Dua akar real sama (kembar)")
add_option(doc, "C", "Tidak memiliki akar real")
add_option(doc, "D", "Satu akar real dan satu akar imajiner")
add_meta(doc, "A", "radio")
add_blank(doc)

# ─── SOAL 8: Jumlah & hasil kali akar (Vieta) ───────────────────────────────
add_question(doc, 8, "Diketahui persamaan kuadrat x² + px + q = 0 mempunyai akar-akar x₁ dan x₂.")
p = doc.add_paragraph("Manakah pernyataan yang benar menurut rumus Vieta? (Pilih semua yang benar)")
p.runs[0].font.size = Pt(11)
add_option(doc, "A", "x₁ + x₂ = -p")
add_option(doc, "B", "x₁ × x₂ = q")
add_option(doc, "C", "x₁ + x₂ = p")
add_option(doc, "D", "x₁ × x₂ = -q")
add_meta(doc, "A, B", "checkbox")
add_blank(doc)

# ─── SOAL 9: Integral ───────────────────────────────────────────────────────
add_question(doc, 9, "Hitunglah hasil dari integral berikut:")
add_math_block(doc, "∫(3x² + 2x - 5) dx")
add_option(doc, "A", "x³ + x² - 5x + C")
add_option(doc, "B", "3x³ + 2x² - 5x + C")
add_option(doc, "C", "x³ + x² + C")
add_option(doc, "D", "6x + 2 + C")
add_meta(doc, "A", "radio")
add_blank(doc)

# ─── SOAL 10: Limit ─────────────────────────────────────────────────────────
add_question(doc, 10, "Tentukan nilai limit berikut:")
add_math_block(doc, "lim(x→2) (x² - 4) / (x - 2)")
add_option(doc, "A", "0")
add_option(doc, "B", "2")
add_option(doc, "C", "4")
add_option(doc, "D", "Tidak terdefinisi")
add_meta(doc, "C", "radio")
add_blank(doc)

# ════════════════════════════════════════════════════════════════════════════
# ─── SOAL CODING 11: Output Python ──────────────────────────────────────────
# ════════════════════════════════════════════════════════════════════════════
add_question(doc, 11, "Perhatikan kode Python berikut:")
add_code_block(doc,
    "x = 5\n"
    "y = 3\n"
    "print(x ** y)"
)
p = doc.add_paragraph("Apakah output dari kode di atas?")
p.runs[0].font.size = Pt(11)
add_option(doc, "A", "15")
add_option(doc, "B", "8")
add_option(doc, "C", "125")
add_option(doc, "D", "Error")
add_meta(doc, "C", "radio")
add_blank(doc)

# ─── SOAL CODING 12: JavaScript async/await ─────────────────────────────────
add_question(doc, 12, "Perhatikan kode JavaScript berikut:")
add_code_block(doc,
    "async function fetchData() {\n"
    "  const res = await fetch('https://api.example.com/data');\n"
    "  const json = await res.json();\n"
    "  return json;\n"
    "}\n"
    "\n"
    "fetchData().then(data => console.log(data));"
)
p = doc.add_paragraph("Manakah pernyataan yang BENAR tentang kode di atas? (Pilih semua yang benar)")
p.runs[0].font.size = Pt(11)
add_option(doc, "A", "fetchData() mengembalikan sebuah Promise")
add_option(doc, "B", "await hanya bisa digunakan di dalam fungsi async")
add_option(doc, "C", "Kode ini akan langsung memblokir thread utama browser")
add_option(doc, "D", "res.json() juga mengembalikan Promise yang perlu di-await")
add_meta(doc, "A, B, D", "checkbox")
add_blank(doc)

# ─── SOAL CODING 13: Big O Notation ─────────────────────────────────────────
add_question(doc, 13, "Perhatikan fungsi berikut:")
add_code_block(doc,
    "function cariDuplikat(arr) {\n"
    "  for (let i = 0; i < arr.length; i++) {\n"
    "    for (let j = i + 1; j < arr.length; j++) {\n"
    "      if (arr[i] === arr[j]) return true;\n"
    "    }\n"
    "  }\n"
    "  return false;\n"
    "}"
)
p = doc.add_paragraph("Berapakah kompleksitas waktu (time complexity) fungsi tersebut?")
p.runs[0].font.size = Pt(11)
add_option(doc, "A", "O(1)")
add_option(doc, "B", "O(n)")
add_option(doc, "C", "O(n²)")
add_option(doc, "D", "O(log n)")
add_meta(doc, "C", "radio")
add_blank(doc)

# ─── SOAL CODING 14: SQL ────────────────────────────────────────────────────
add_question(doc, 14, "Perhatikan query SQL berikut:")
add_code_block(doc,
    "SELECT u.name, COUNT(o.id) AS total_order\n"
    "FROM users u\n"
    "LEFT JOIN orders o ON u.id = o.user_id\n"
    "WHERE o.status = 'completed'\n"
    "GROUP BY u.name\n"
    "HAVING COUNT(o.id) > 5\n"
    "ORDER BY total_order DESC;"
)
p = doc.add_paragraph("Apa yang dilakukan query tersebut?")
p.runs[0].font.size = Pt(11)
add_option(doc, "A", "Menampilkan semua user beserta jumlah order mereka")
add_option(doc, "B", "Menampilkan user yang memiliki lebih dari 5 order berstatus 'completed', diurutkan dari terbanyak")
add_option(doc, "C", "Menampilkan user tanpa order sama sekali")
add_option(doc, "D", "Menghapus order yang statusnya bukan 'completed'")
add_meta(doc, "B", "radio")
add_blank(doc)

# ─── SOAL CODING 15: Bug Hunting ────────────────────────────────────────────
add_question(doc, 15, "Kode berikut seharusnya menghitung faktorial dari n, tapi ada bug. Temukan baris yang salah!")
add_code_block(doc,
    "function faktorial(n) {\n"
    "  if (n === 0) return 1;\n"
    "  return n * faktorial(n);   // baris ini\n"
    "}"
)
p = doc.add_paragraph("Apa yang menyebabkan bug pada kode di atas?")
p.runs[0].font.size = Pt(11)
add_option(doc, "A", "Base case salah, seharusnya n === 1")
add_option(doc, "B", "Rekursi memanggil faktorial(n) bukan faktorial(n - 1) sehingga infinite loop")
add_option(doc, "C", "Fungsi tidak mengembalikan nilai apapun")
add_option(doc, "D", "Operator * tidak bisa digunakan untuk rekursi")
add_meta(doc, "B", "radio")
add_blank(doc)

# ─── Simpan file ─────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), "soal-template.docx")
doc.save(output_path)
print(f"✅ File berhasil dibuat: {output_path}")
print(f"   Total soal: 15 (5 lama + 5 rumus matematika + 5 coding)")
